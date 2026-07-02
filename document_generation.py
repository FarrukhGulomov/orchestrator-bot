"""
Formatted document deliverables (Word + PDF).

Some requests shouldn't get a chat reply at all — they should get a polished,
shareable FILE. The clearest example: "tijorat taklifi" (commercial
proposal) / cost-and-resource estimate questions. Dumping that as raw chat
text defeats the point; the user wants something they can forward to a
client or a center director as-is.

Flow:
  1. generate_proposal_content() asks the routed agent's own model for the
     document's CONTENT as strict JSON (title, intro, sections, an optional
     cost/resource table, closing) — not prose to parse, a real schema.
  2. render_docx() / render_pdf() turn that JSON into an actually nice-looking
     file: styled title, real headings, a real table for cost/resource
     breakdowns, a footer. Both renderers consume the exact same content, so
     the two files always match.

Content JSON shape (all fields produced by the LLM, see _SCHEMA_INSTRUCTION):
{
  "title": str,
  "subtitle": str | None,
  "intro": str,
  "sections": [{"heading": str, "body": str}, ...],
  "table": {"title": str | None, "headers": [str, ...], "rows": [[str,...],...]} | None,
  "closing": str,
}
"""

import io
import json
import logging
from datetime import date

from agents import Agent
from config import settings
from llm_clients import claude_generate_json
from router import model_for

logger = logging.getLogger(__name__)

_SCHEMA_INSTRUCTION = """
DOCUMENT MODE.
You are generating the CONTENT for a polished, professional document (e.g. a
commercial proposal / resource-and-cost estimate, or a formal report) based
on the user's request.

LANGUAGE: detect from the GRAMMAR of the request, NOT from BA/IT keywords.
Uzbek grammar markers: "faylda", "tayyorla", "qiling", "menga", "kerak", "-da", "-ni".
Russian grammar markers: "в", "на", "для", "мне", "нужно", "сделай".
DEFAULT to UZBEK when unsure. Keep BA/IT terms (BRD, KPI, SQL, etc.) in English.

Be concrete: real numbers/estimates where the request gives enough to infer
them (briefly state assumptions if you must estimate), real deliverables,
real timeline — not vague generalities. This becomes a real file someone may
forward to a client, so it must stand on its own without any chat context.

Respond with ONLY a JSON object matching exactly this shape — no prose, no
markdown code fences, nothing outside the JSON:
{"title": "string", "subtitle": "string or null", "intro": "string",
 "sections": [{"heading": "string", "body": "string"}],
 "table": {"title": "string or null", "headers": ["string", ...], "rows": [["string", ...], ...]} or null,
 "closing": "string"}

Use "table" for any cost/resource/timeline breakdown — that is usually the
whole point of this document. Use 3-6 sections. Keep each section's body
concise (this is a document people skim, not an essay). Set "table" to null
if there is genuinely nothing tabular to show.
"""

_FOOTER_TEXT = "Master Orchestrator AI Team tomonidan tayyorlandi"


async def generate_proposal_content(agent: Agent, user_text: str) -> dict:
    system_prompt = agent.system + "\n" + _SCHEMA_INSTRUCTION
    # Use the agent's own MAIN model with a full token budget — the default
    # claude_generate_json settings (fast model, 512 tokens) are for routing
    # and would truncate a real document's JSON mid-object.
    model, _label = model_for(agent, "high")
    raw = await claude_generate_json(
        system_prompt,
        [{"role": "user", "content": user_text}],
        model=model,
        max_tokens=settings.max_output_tokens,
    )
    try:
        data = json.loads(raw)
        if not isinstance(data, dict):
            raise ValueError(f"expected JSON object, got {type(data).__name__}")
    except (json.JSONDecodeError, ValueError) as exc:
        # Model returned malformed/truncated JSON — degrade gracefully instead
        # of crashing the whole /proposal flow with an uncaught exception.
        logger.warning(
            "Proposal JSON parse failed (%s); raw head: %.200s", exc, raw
        )
        data = {
            "title": "Taklif",
            "subtitle": None,
            "intro": raw.strip()[:3000] if raw.strip() else (
                "Hujjat tarkibini tayyorlab bo'lmadi. Iltimos, so'rovni qisqartirib "
                "yoki aniqlashtirib qaytadan yuboring."
            ),
            "sections": [],
            "table": None,
            "closing": "",
        }
    data.setdefault("title", "Taklif")
    data.setdefault("subtitle", None)
    data.setdefault("intro", "")
    data.setdefault("sections", [])
    data.setdefault("table", None)
    data.setdefault("closing", "")
    return data


def render_docx(data: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    doc.add_heading(data.get("title") or "Taklif", level=0)

    if data.get("subtitle"):
        p = doc.add_paragraph(data["subtitle"])
        if p.runs:
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(11)

    meta = doc.add_paragraph(f"Sana: {date.today().strftime('%d.%m.%Y')}")
    if meta.runs:
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    if data.get("intro"):
        doc.add_paragraph(data["intro"])

    for section in data.get("sections") or []:
        doc.add_heading(str(section.get("heading", "")), level=2)
        body = str(section.get("body", ""))
        for para in body.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    table_data = data.get("table")
    if table_data and table_data.get("headers") and table_data.get("rows"):
        if table_data.get("title"):
            doc.add_heading(str(table_data["title"]), level=2)
        headers = table_data["headers"]
        rows = table_data["rows"]
        table = doc.add_table(rows=1, cols=len(headers))
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass  # template without this built-in style — keep the default look
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            for p in hdr_cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                if i < len(cells):
                    cells[i].text = str(val)

    if data.get("closing"):
        doc.add_paragraph()
        doc.add_paragraph(data["closing"])

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_run = footer_p.add_run(_FOOTER_TEXT)
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _esc(text: str) -> str:
    # reportlab's Paragraph uses a small XML-like markup; escape user/LLM text
    # so stray '<' or '&' characters don't break rendering.
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def render_pdf(data: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], fontSize=20, spaceAfter=6)
    subtitle_style = ParagraphStyle(
        "SubtitleX", parent=styles["Normal"], fontSize=11, textColor=colors.grey, spaceAfter=2
    )
    meta_style = ParagraphStyle(
        "MetaX", parent=styles["Normal"], fontSize=9, textColor=colors.grey, spaceAfter=12
    )
    heading_style = ParagraphStyle(
        "HeadingX", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6
    )
    body_style = styles["Normal"]
    footer_style = ParagraphStyle("FooterX", parent=styles["Normal"], fontSize=8, textColor=colors.grey)

    flow = [Paragraph(_esc(data.get("title") or "Taklif"), title_style)]
    if data.get("subtitle"):
        flow.append(Paragraph(_esc(data["subtitle"]), subtitle_style))
    flow.append(Paragraph(f"Sana: {date.today().strftime('%d.%m.%Y')}", meta_style))

    if data.get("intro"):
        flow.append(Paragraph(_esc(data["intro"]), body_style))
        flow.append(Spacer(1, 8))

    for section in data.get("sections") or []:
        flow.append(Paragraph(_esc(str(section.get("heading", ""))), heading_style))
        for para in str(section.get("body", "")).split("\n"):
            if para.strip():
                flow.append(Paragraph(_esc(para.strip()), body_style))
                flow.append(Spacer(1, 4))

    table_data = data.get("table")
    if table_data and table_data.get("headers") and table_data.get("rows"):
        if table_data.get("title"):
            flow.append(Paragraph(_esc(str(table_data["title"])), heading_style))
        headers = [str(h) for h in table_data["headers"]]
        rows = [[str(v) for v in row] for row in table_data["rows"]]
        t = Table([headers] + rows, hAlign="LEFT")
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2D3142")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F5F5")]),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        flow.append(Spacer(1, 6))
        flow.append(t)
        flow.append(Spacer(1, 10))

    if data.get("closing"):
        flow.append(Paragraph(_esc(data["closing"]), body_style))

    flow.append(Spacer(1, 20))
    flow.append(Paragraph(_FOOTER_TEXT, footer_style))

    doc.build(flow)
    return buf.getvalue()
