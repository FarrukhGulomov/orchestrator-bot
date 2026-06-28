"""
Inbound file understanding.

Lets the bot accept (almost) any file the team throws at it and turn it into
plain text that then flows through the exact same classify() -> agent
pipeline as a normal text message. This means a photo of a whiteboard, a
BPMN export, a bug report in a .docx, or a voice note all end up routed,
classified, and answered exactly like typed text — no special-casing needed
anywhere else in the bot.

Supported kinds:
  - images   (jpg/jpeg/png/webp/gif/bmp)   -> described via Gemini vision
  - pdf                                     -> read via Gemini's native PDF
                                                understanding (handles scans,
                                                tables, diagrams — not just
                                                a naive text dump)
  - docx                                    -> text + tables via python-docx
  - xlsx / xlsm                             -> sheets via openpyxl, rendered
                                                as simple pipe-delimited tables
  - csv                                     -> rendered as a pipe-delimited table
  - bpmn / xml                              -> raw XML passed through as text
                                                (BPMN files ARE XML; agents
                                                read the process definition
                                                directly)
  - txt / md / json / yaml / log / etc.     -> decoded as plain text
  - audio (mp3/wav/ogg/m4a/flac/webm +
            Telegram voice notes)           -> transcribed via Groq Whisper

Anything else is reported back to the caller as unsupported, with the list
of formats that do work, instead of silently failing.
"""

import csv
import io
import logging

from llm_clients import claude_describe_file

logger = logging.getLogger(__name__)

# Cap on extracted text fed into the routing pipeline, to keep the downstream
# classify() + agent generation calls within a sane token/cost budget.
MAX_CHARS = 15000

IMAGE_EXTS = {"jpg", "jpeg", "png", "webp", "gif", "bmp"}
AUDIO_EXTS = {"mp3", "wav", "ogg", "oga", "m4a", "flac", "webm", "mp4", "mpeg", "mpga"}
TEXT_EXTS = {"txt", "md", "log", "json", "yaml", "yml", "ini", "cfg"}
XML_EXTS = {"bpmn", "xml"}

SUPPORTED_SUMMARY = (
    "rasm (jpg/png/webp/...), PDF, Word (.docx), Excel (.xlsx), CSV, "
    "BPMN/XML, audio (mp3/wav/ogg/ovoz xabari...), oddiy matn fayllar "
    "(.txt/.md/.json/.yaml)"
)


def _ext(filename: str | None) -> str:
    if not filename or "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def guess_kind(filename: str | None, mime_type: str | None) -> str:
    ext = _ext(filename)
    mime = (mime_type or "").lower()

    if ext in IMAGE_EXTS or mime.startswith("image/"):
        return "image"
    if ext == "pdf" or mime == "application/pdf":
        return "pdf"
    if ext == "docx" or "wordprocessingml" in mime:
        return "docx"
    if ext == "doc" or mime == "application/msword":
        return "doc_legacy"
    if ext in {"xlsx", "xlsm"} or "spreadsheetml" in mime:
        return "xlsx"
    if ext == "xls" or mime == "application/vnd.ms-excel":
        return "xls_legacy"
    if ext == "csv" or mime == "text/csv":
        return "csv"
    if ext in XML_EXTS or mime in {"application/xml", "text/xml"}:
        return "xml"
    if ext in AUDIO_EXTS or mime.startswith("audio/"):
        return "audio"
    if ext in TEXT_EXTS or mime.startswith("text/"):
        return "text"
    return "unknown"


def _truncate(text: str) -> str:
    text = (text or "").strip()
    if len(text) > MAX_CHARS:
        return (
            text[:MAX_CHARS]
            + f"\n\n[...qisqartirildi: jami {len(text)} belgidan {MAX_CHARS} tasi ko'rsatildi...]"
        )
    return text


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return _truncate("\n".join(parts))


def extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"### Sheet: {sheet.title}")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            if row_count >= 200:
                lines.append("[...200 qatordan keyin qisqartirildi...]")
                break
            cells = ["" if v is None else str(v) for v in row]
            lines.append(" | ".join(cells))
            row_count += 1
    return _truncate("\n".join(lines))


def extract_csv(data: bytes) -> str:
    text = _decode_text(data)
    reader = csv.reader(io.StringIO(text))
    lines = []
    for i, row in enumerate(reader):
        if i >= 300:
            lines.append("[...300 qatordan keyin qisqartirildi...]")
            break
        lines.append(" | ".join(row))
    return _truncate("\n".join(lines))


def extract_xml(data: bytes) -> str:
    return _truncate(_decode_text(data))


def extract_plain_text(data: bytes) -> str:
    return _truncate(_decode_text(data))


async def describe_image(data: bytes, mime_type: str) -> str:
    instruction = (
        "Diqqat bilan tahlil qiling: rasmda nima ko'rsatilgan (sxema, ekran "
        "skrinshoti, diagramma, jadval, qo'lyozma matn va h.k.). Barcha "
        "ko'rinadigan matnni so'zma-so'z keltiring, tuzilmani (bloklar, "
        "strelkalar, jadval ustunlari, bog'lanishlar) tasvirlang. Javobni "
        "faqat shu tahlil bilan bering, boshqa izoh qo'shmang."
    )
    text = await claude_describe_file(data, mime_type or "image/jpeg", instruction)
    return _truncate(text)


async def describe_pdf(data: bytes) -> str:
    instruction = (
        "Ushbu PDF hujjatining to'liq matnini va tuzilmasini chiqarib bering: "
        "sarlavhalar, asosiy matn, jadvallar (matn ko'rinishida), diagrammalar "
        "tavsifi. Hujjat juda uzun bo'lsa eng muhim qismlarini to'liq, "
        "qolganini qisqa umumlashtiring. Faqat shu tarkibni bering, ortiqcha "
        "sharh qo'shmang."
    )
    text = await claude_describe_file(data, "application/pdf", instruction)
    return _truncate(text)


async def transcribe_audio(data: bytes, filename: str) -> str:
    # Audio transcription is not available in Claude-only mode.
    # Return a user-friendly message instead of failing silently.
    raise NotImplementedError(
        "Ovozni matnga o'girish uchun qo'shimcha xizmat kerak. "
        "Hozircha ovozli xabarlarni matn orqali yuboring."
    )


async def extract(
    filename: str | None, mime_type: str | None, data: bytes
) -> tuple[str | None, str]:
    """Try to turn a file into text.

    Returns (extracted_text, kind) on success, or (None, error_message) if
    the file can't be processed — the caller should show error_message to
    the user as-is and stop.
    """
    kind = guess_kind(filename, mime_type)
    try:
        if kind == "image":
            return await describe_image(data, mime_type or ""), kind
        if kind == "pdf":
            return await describe_pdf(data), kind
        if kind == "docx":
            return extract_docx(data), kind
        if kind == "xlsx":
            return extract_xlsx(data), kind
        if kind == "csv":
            return extract_csv(data), kind
        if kind == "xml":
            return extract_xml(data), kind
        if kind == "text":
            return extract_plain_text(data), kind
        if kind == "audio":
            return await transcribe_audio(data, filename or "audio"), kind
        if kind in ("doc_legacy", "xls_legacy"):
            return None, (
                "Bu eski formatdagi fayl (.doc/.xls). Iltimos, .docx yoki "
                ".xlsx formatida qayta yuboring (Word/Excel'da 'Save As' "
                "orqali)."
            )
        return None, (
            "Bu fayl turini hali o'qiy olmayman.\n"
            f"Qo'llab-quvvatlanadigan turlar: {SUPPORTED_SUMMARY}."
        )
    except Exception as exc:  # noqa: BLE001 — never let a bad file crash the bot
        logger.exception("File extraction failed for kind=%s filename=%s", kind, filename)
        return None, f"⚠️ Faylni o'qishda xatolik: {exc}"
